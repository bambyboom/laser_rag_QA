from __future__ import annotations

import logging
from pathlib import Path
from typing import TYPE_CHECKING

from pypdf import PdfReader
from pypdf.errors import PdfReadError
from docx import Document as DocxDocument

if TYPE_CHECKING:
    from config import Settings

logger = logging.getLogger("laser-rag.loader")

# ---------------------------------------------------------------------------
#  encoding fallback chain for plain-text files
# ---------------------------------------------------------------------------
_ENCODINGS: tuple[str, ...] = ("utf-8", "gbk", "gb18030", "latin-1")


class DocumentLoader:
    """Load documents from a directory and return uniform text records.

    Supported formats: PDF, DOCX, TXT (configurable via Settings).

    Usage::

        loader = DocumentLoader()
        docs = loader.load_all()
        # docs → [{"filename": "manual.pdf", "text": "..."}, ...]
    """

    def __init__(self, raw_dir: Path | str | None = None) -> None:
        from config import settings

        self._settings: Settings = settings
        self._raw_dir = Path(raw_dir) if raw_dir else (settings.data_dir / "raw")
        self._raw_dir.mkdir(parents=True, exist_ok=True)

    # ------------------------------------------------------------------
    #  public API
    # ------------------------------------------------------------------
    def load_all(self) -> list[dict[str, str]]:
        """Walk ``raw_dir``, load every supported file, return records."""
        records: list[dict[str, str]] = []
        if not self._raw_dir.exists():
            logger.warning("Raw directory does not exist: %s", self._raw_dir)
            return records

        for file_path in sorted(self._raw_dir.iterdir()):
            if not file_path.is_file():
                continue

            suffix = file_path.suffix.lower()
            if suffix not in self._settings.supported_suffixes:
                logger.debug("Skipping unsupported file: %s", file_path.name)
                continue

            try:
                text = self._load_one(file_path)
            except Exception:
                logger.exception("Failed to load %s — skipping", file_path.name)
                continue

            if not text.strip():
                # PDFs may be scanned (no text layer) — keep them for OCR.
                # Other formats with empty content are genuinely empty → skip.
                if suffix == ".pdf":
                    logger.warning(
                        "Empty content in %s — keeping for OCR processing",
                        file_path.name,
                    )
                else:
                    logger.warning(
                        "Empty content in %s — skipping", file_path.name
                    )
                    continue

            records.append({"filename": file_path.name, "text": text})
            logger.info("Loaded %s (%d chars)", file_path.name, len(text))

        logger.info("Loaded %d document(s) from %s", len(records), self._raw_dir)
        return records

    def load_file(self, file_path: Path | str) -> dict[str, str]:
        """Load a single file, return a single record."""
        fp = Path(file_path)
        text = self._load_one(fp)
        return {"filename": fp.name, "text": text}

    # ------------------------------------------------------------------
    #  internal dispatcher
    # ------------------------------------------------------------------
    def _load_one(self, file_path: Path) -> str:
        suffix = file_path.suffix.lower()

        if suffix == ".pdf":
            return self._load_pdf(file_path)
        elif suffix in (".docx", ".doc"):
            return self._load_docx(file_path)
        elif suffix in (".txt", ".md", ".html"):
            return self._load_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

    # ------------------------------------------------------------------
    #  per-format parsers
    # ------------------------------------------------------------------
    @staticmethod
    def _load_pdf(fp: Path) -> str:
        """Extract all page text from a PDF."""
        try:
            reader = PdfReader(str(fp))
        except PdfReadError as exc:
            raise RuntimeError(f"Corrupted or unreadable PDF: {fp}") from exc

        pages: list[str] = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                pages.append(text)
            else:
                logger.debug("No extractable text on page %d of %s", i + 1, fp.name)
        return "\n".join(pages)

    @staticmethod
    def _load_docx(fp: Path) -> str:
        """Extract all paragraph text from a DOCX file."""
        try:
            doc = DocxDocument(str(fp))
        except Exception as exc:
            raise RuntimeError(f"Failed to open DOCX: {fp}") from exc

        paragraphs: list[str] = []
        for para in doc.paragraphs:
            text = para.text
            if text.strip():
                paragraphs.append(text)

        # also try tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        return "\n".join(paragraphs)

    @staticmethod
    def _load_text(fp: Path) -> str:
        """Read a plain-text file with encoding fallback."""
        for enc in _ENCODINGS:
            try:
                return fp.read_text(encoding=enc)
            except (UnicodeDecodeError, UnicodeError):
                logger.debug("Encoding %s failed for %s", enc, fp.name)
                continue
        raise RuntimeError(
            f"Unable to decode {fp} with any of: {', '.join(_ENCODINGS)}"
        )
